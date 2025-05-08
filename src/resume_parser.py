# src/resume_parser.py
import os
from docx import Document
import PyPDF2

def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
        return None

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    try:
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            full_text = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                full_text.append(page.extract_text())
            return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
        return None

def parse_resume(file_path):
    """
    Parses a resume file (DOCX or PDF) and extracts raw text.
    More structured parsing can be added later.
    """
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return None

    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    text_content = None
    if file_extension == ".docx":
        text_content = extract_text_from_docx(file_path)
    elif file_extension == ".pdf":
        text_content = extract_text_from_pdf(file_path)
    else:
        print(f"Error: Unsupported file type: {file_extension}. Only .docx and .pdf are supported.")
        return None

    if text_content:
        # Basic structuring attempt (can be significantly improved with NLP/regex)
        # For now, just returning raw text.
        # Future: Extract sections like contact, experience, education, skills.
        # Example:
        # contact_info = extract_contact_info(text_content)
        # experience = extract_experience(text_content)
        # return {"raw_text": text_content, "contact": contact_info, ...}
        return {"raw_text": text_content}
    return None

if __name__ == "__main__":
    # Example Usage (Create dummy files for testing)
    # Create a dummy docx
    try:
        from docx import Document as DocWriter
        doc = DocWriter()
        doc.add_paragraph("This is a test DOCX resume.")
        doc.add_paragraph("John Doe - john.doe@example.com - 123-456-7890")
        doc.add_paragraph("Experience: Software Engineer at Tech Corp")
        dummy_docx_path = "dummy_resume.docx"
        doc.save(dummy_docx_path)
        print(f"Created dummy file: {dummy_docx_path}")

        parsed_docx = parse_resume(dummy_docx_path)
        if parsed_docx:
            print("\n--- Parsed DOCX ---")
            print(parsed_docx["raw_text"][:200] + "...") # Print first 200 chars
        os.remove(dummy_docx_path) # Clean up
    except ImportError:
        print("Skipping DOCX example: python-docx not fully available for writing in this environment.")
    except Exception as e:
        print(f"Error in DOCX example: {e}")


    # Note: Creating a dummy PDF programmatically is more complex and often requires
    # libraries like reportlab, which might not be available.
    # For PDF testing, you'd typically use an existing PDF file.
    print("\nFor PDF testing, please place a PDF file in the same directory and update path.")
    # dummy_pdf_path = "path_to_your_dummy_resume.pdf"
    # if os.path.exists(dummy_pdf_path):
    #     parsed_pdf = parse_resume(dummy_pdf_path)
    #     if parsed_pdf:
    #         print("\n--- Parsed PDF ---")
    #         print(parsed_pdf["raw_text"][:200] + "...")
    # else:
    #     print(f"Skipping PDF example: {dummy_pdf_path} not found.")